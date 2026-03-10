"""
CV Generator - produces ATS-friendly DOCX and PDF files.

ATS optimization rules:
- No tables, text boxes, or floating elements
- Standard heading styles (Heading 1, Heading 2)
- Simple paragraphs with clear formatting
- Standard fonts (Calibri)
- No images/logos in DOCX (ATS can't parse them)
- Clean section headers
- Consistent date formatting
- Single column layout
- Hyperlinks use standard OOXML elements (ATS-compatible)
"""

import os
import re
import shutil
import subprocess
import tempfile

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import parse_xml

DOCX_TRANSLATIONS = {
    'en': {
        'summary': 'SUMMARY', 'experience': 'WORK EXPERIENCE', 'skills': 'SKILLS',
        'projects': 'PROJECTS', 'courses': 'COURSES & TRAINING', 'education': 'EDUCATION',
        'languages': 'LANGUAGES', 'certifications': 'CERTIFICATIONS',
    },
    'pl': {
        'summary': 'PODSUMOWANIE', 'experience': 'DOŚWIADCZENIE ZAWODOWE', 'skills': 'UMIEJĘTNOŚCI',
        'projects': 'PROJEKTY', 'courses': 'KURSY I SZKOLENIA', 'education': 'EDUKACJA',
        'languages': 'JĘZYKI', 'certifications': 'CERTYFIKATY',
    },
    'de': {
        'summary': 'ZUSAMMENFASSUNG', 'experience': 'BERUFSERFAHRUNG', 'skills': 'FÄHIGKEITEN',
        'projects': 'PROJEKTE', 'courses': 'KURSE & WEITERBILDUNG', 'education': 'AUSBILDUNG',
        'languages': 'SPRACHEN', 'certifications': 'ZERTIFIZIERUNGEN',
    },
    'fr': {
        'summary': 'RÉSUMÉ', 'experience': 'EXPÉRIENCE PROFESSIONNELLE', 'skills': 'COMPÉTENCES',
        'projects': 'PROJETS', 'courses': 'FORMATIONS', 'education': 'FORMATION',
        'languages': 'LANGUES', 'certifications': 'CERTIFICATIONS',
    },
    'es': {
        'summary': 'RESUMEN', 'experience': 'EXPERIENCIA LABORAL', 'skills': 'HABILIDADES',
        'projects': 'PROYECTOS', 'courses': 'CURSOS Y FORMACIÓN', 'education': 'EDUCACIÓN',
        'languages': 'IDIOMAS', 'certifications': 'CERTIFICACIONES',
    },
}

def _t(lang, key):
    """Translate section title."""
    translations = DOCX_TRANSLATIONS.get(lang, DOCX_TRANSLATIONS['en'])
    return translations.get(key, DOCX_TRANSLATIONS['en'].get(key, key))

def _setup_styles(doc):
    """Configure document styles for ATS-friendly output."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    if 'CVHeading1' not in [s.name for s in doc.styles]:
        h1 = doc.styles.add_style('CVHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1.base_style = doc.styles['Normal']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(13)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1A, 0x23, 0x3B)
        h1.paragraph_format.space_before = Pt(16)
        h1.paragraph_format.space_after = Pt(6)
        h1.paragraph_format.keep_with_next = True

    if 'CVHeading2' not in [s.name for s in doc.styles]:
        h2 = doc.styles.add_style('CVHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2.base_style = doc.styles['Normal']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(11)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x1A, 0x23, 0x3B)
        h2.paragraph_format.space_before = Pt(8)
        h2.paragraph_format.space_after = Pt(1)
        h2.paragraph_format.keep_with_next = True

def _add_horizontal_line(doc):
    """Add a subtle horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

def _add_hyperlink(paragraph, text, url, font_size=None, font_color=None, bold=False):
    """Add a clickable hyperlink to a paragraph (ATS-compatible OOXML element)."""
    url = _sanitize_url(url)
    if not url:
        run = paragraph.add_run(text)
        if font_size:
            run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        f' r:id="{r_id}"/>'
    )
    run_elem = parse_xml(
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:rPr></w:rPr></w:r>'
    )
    rPr = run_elem.find(qn('w:rPr'))
    rFonts = parse_xml(
        '<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:ascii="Calibri" w:hAnsi="Calibri"/>'
    )
    rPr.append(rFonts)
    if font_size:
        sz = parse_xml(
            f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:val="{int(font_size * 2)}"/>'
        )
        rPr.append(sz)
    if font_color:
        color = parse_xml(
            f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:val="{font_color}"/>'
        )
        rPr.append(color)
    if bold:
        b = parse_xml(
            '<w:b xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        rPr.append(b)
    u = parse_xml(
        '<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:val="single"/>'
    )
    rPr.append(u)
    t = parse_xml(
        f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' xml:space="preserve">{_escape_xml(text)}</w:t>'
    )
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)

def _escape_xml(text):
    """Escape XML special characters."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

def _html_to_plain_text(html):
    """Strip HTML tags and convert to plain text for ATS-friendly output."""
    if not html:
        return ''
    text = re.sub(r'<br\s*/?>', '\n', html, flags=re.IGNORECASE)
    text = re.sub(r'</(?:p|div|li|h[1-6])>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('&nbsp;', ' ').replace('&quot;', '"')
    lines = [line.strip() for line in text.split('\n')]
    return '\n'.join(line for line in lines if line)

def _sanitize_url(url):
    """Ensure URL has a safe scheme (http/https only)."""
    if not url:
        return ''
    url = url.strip()
    if url.lower().startswith(('javascript:', 'data:', 'vbscript:')):
        return ''
    if not url.startswith(('http://', 'https://', 'mailto:', 'tel:')):
        url = f'https://{url}'
    return url

def _contact_href(contact):
    """Generate appropriate URL for a contact based on type."""
    ctype = contact.get('type', '')
    value = contact.get('value', '')
    if ctype == 'email':
        return f'mailto:{value}'
    elif ctype == 'phone':
        return f'tel:{value.replace(" ", "")}'
    elif ctype in ('website', 'linkedin', 'github', 'twitter'):
        return _sanitize_url(value)
    return ''

def _add_name_header(doc, personal):
    """Add name and title at the top."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(personal['name'])
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x3B)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(personal['title'])
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x65, 0x80)
    run.font.name = 'Calibri'

    contacts = personal.get('contacts', [])
    if not contacts:
        contact_parts = []
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(' | '.join(contact_parts))
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x55, 0x65, 0x80)
        run.font.name = 'Calibri'
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        first = True
        for contact in contacts:
            value = contact.get('value', '')
            if not value:
                continue
            if not first:
                run = p.add_run(' | ')
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x55, 0x65, 0x80)
                run.font.name = 'Calibri'
            first = False
            href = _contact_href(contact) if contact.get('link') else ''
            if href:
                _add_hyperlink(p, value, href, font_size=9.5, font_color='556580')
            else:
                run = p.add_run(value)
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x55, 0x65, 0x80)
                run.font.name = 'Calibri'

    _add_horizontal_line(doc)

def _add_clause(doc, clause_text):
    """Add GDPR clause at the bottom."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clause_text)
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    run.font.name = 'Calibri'

def _add_summary(doc, summary, lang='en'):
    """Add professional summary section."""
    p = doc.add_paragraph(_t(lang, 'summary'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(4)

def _add_work_experience(doc, data, lang='en'):
    """Add work experience section - supports both old and new grouped format."""
    p = doc.add_paragraph(_t(lang, 'experience'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True

    groups = data.get('employer_groups', [])
    if groups:
        first = True
        for group in groups:
            if group.get('hidden'):
                continue
            group_url = group.get('url', '')
            for pos in group.get('positions', []):
                company_name = pos.get('display_company', group['group_name'])
                p = doc.add_paragraph(style='CVHeading2')
                if group_url:
                    _add_hyperlink(p, company_name, group_url,
                                   font_size=11, font_color='1A233B', bold=True)
                else:
                    run = p.add_run(company_name)
                p.paragraph_format.space_before = Pt(2) if first else Pt(8)
                p.paragraph_format.keep_with_next = True
                first = False

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(pos['role'])
                run.font.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                date_str = f"{pos['date_from']} - {pos['date_to']}"
                run = p.add_run(f"    |    {date_str}")
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                desc_format = pos.get('desc_format', 'bullets')
                if desc_format in ('rich', 'paragraph') and pos.get('rich_description'):
                    plain = _html_to_plain_text(pos['rich_description'])
                    if plain:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.left_indent = Inches(0.2)
                        run = p.add_run(plain)
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                else:
                    bullets = pos.get('bullets', [])
                    for bi, bullet in enumerate(bullets):
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.left_indent = Inches(0.2)
                        if bi < len(bullets) - 1:
                            p.paragraph_format.keep_with_next = True
                        run = p.add_run(f"• {bullet}")
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        for i, exp in enumerate(data.get('work_experience', [])):
            if exp.get('hidden'):
                continue
            p = doc.add_paragraph(style='CVHeading2')
            run = p.add_run(exp['company'])
            p.paragraph_format.space_before = Pt(8) if i > 0 else Pt(2)
            p.paragraph_format.keep_with_next = True

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(exp['role'])
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            date_str = f"{exp['date_from']} - {exp['date_to']}"
            run = p.add_run(f"    |    {date_str}")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

            bullets = exp.get('bullets', [])
            if not bullets and exp.get('description'):
                bullets = [exp['description']]
            for bi, bullet in enumerate(bullets):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.2)
                if bi < len(bullets) - 1:
                    p.paragraph_format.keep_with_next = True
                run = p.add_run(f"• {bullet}")
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def _add_education(doc, education, lang='en'):
    """Add education section."""
    p = doc.add_paragraph(_t(lang, 'education'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True

    for edu in education:
        p = doc.add_paragraph(style='CVHeading2')
        edu_url = edu.get('url', '')
        if edu_url:
            _add_hyperlink(p, edu['institution'], edu_url,
                           font_size=11, font_color='1A233B', bold=True)
        else:
            run = p.add_run(edu['institution'])

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        degree_text = ' — '.join(filter(None, [edu.get('level', ''), edu.get('degree', '')]))
        run = p.add_run(degree_text)
        run.font.size = Pt(10)

        date_str = f"{edu['date_from']} - {edu['date_to']}"
        run = p.add_run(f"    |    {date_str}")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

def _add_skills(doc, skills, lang='en'):
    """Add skills section - categories with skill tags."""
    p = doc.add_paragraph(_t(lang, 'skills'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True

    for si, cat in enumerate(skills):
        category = cat.get('category', '')
        items = [item for item in cat.get('items', []) if item]
        if not category and not items:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        if si < len(skills) - 1:
            p.paragraph_format.keep_with_next = True
        if category:
            run = p.add_run(f"{category}: ")
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x3B)
        run = p.add_run(', '.join(items))
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def _add_projects(doc, projects, lang='en'):
    """Add projects section."""
    p = doc.add_paragraph(_t(lang, 'projects'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True

    for pi, proj in enumerate(projects):
        name = proj.get('name', '')
        if not name:
            continue
        url = proj.get('url', '')
        p = doc.add_paragraph(style='CVHeading2')
        p.paragraph_format.keep_with_next = True
        if url:
            _add_hyperlink(p, name, url if url.startswith('http') else f'https://{url}',
                           font_size=11, font_color='1A233B', bold=True)
        else:
            run = p.add_run(name)
        p.paragraph_format.space_before = Pt(2) if pi == 0 else Pt(8)

        role = proj.get('role', '')
        date_from = proj.get('date_from', '')
        date_to = proj.get('date_to', '')
        if role or date_from:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            if role:
                run = p.add_run(role)
                run.font.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            if date_from or date_to:
                sep = '    |    ' if role else ''
                run = p.add_run(f"{sep}{date_from} - {date_to}")
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        desc = proj.get('description', '')
        if desc:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(desc)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def _add_courses(doc, courses, lang='en'):
    """Add courses/training section."""
    p = doc.add_paragraph(_t(lang, 'courses'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True

    for ci, course in enumerate(courses):
        name = course.get('name', '')
        if not name:
            continue
        url = course.get('url', '')
        provider = course.get('provider', '')
        date = course.get('date', '')

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        if ci < len(courses) - 1:
            p.paragraph_format.keep_with_next = True

        run = p.add_run('• ')
        run.font.size = Pt(9.5)
        if url:
            _add_hyperlink(p, name, url if url.startswith('http') else f'https://{url}',
                           font_size=9.5, font_color='333333')
        else:
            run = p.add_run(name)
            run.font.size = Pt(9.5)
            run.font.bold = True
        if provider:
            run = p.add_run(f' — {provider}')
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if date:
            run = p.add_run(f'  ({date})')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

def _add_languages(doc, languages, lang='en'):
    """Add languages section."""
    p = doc.add_paragraph(_t(lang, 'languages'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True

    for entry in languages:
        p = doc.add_paragraph()
        run = p.add_run(f"{entry['language']}")
        run.font.bold = True
        run = p.add_run(f" - {entry['level']}")

def _add_certifications(doc, certifications, lang='en'):
    """Add certifications section."""
    p = doc.add_paragraph(_t(lang, 'certifications'), style='CVHeading1')
    p.paragraph_format.keep_with_next = True

    for cert_group in certifications:
        p = doc.add_paragraph(style='CVHeading2')
        p.paragraph_format.keep_with_next = True
        issuer_url = cert_group.get('issuer_url', '')
        if issuer_url:
            _add_hyperlink(p, cert_group['issuer'], issuer_url,
                           font_size=11, font_color='1A233B', bold=True)
        else:
            run = p.add_run(cert_group['issuer'])

        items = cert_group['items']
        for ii, item in enumerate(items):
            if isinstance(item, str):
                item_name = item
                item_url = ''
            else:
                item_name = item.get('name', '')
                item_url = item.get('url', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.2)
            if ii < len(items) - 1:
                p.paragraph_format.keep_with_next = True
            if item_url:
                run = p.add_run('• ')
                run.font.size = Pt(9.5)
                _add_hyperlink(p, item_name, item_url, font_size=9.5, font_color='333333')
            else:
                run = p.add_run(f"• {item_name}")
                run.font.size = Pt(9.5)

def generate_docx(data):
    """Generate ATS-friendly DOCX file."""
    if not data or not isinstance(data, dict):
        raise ValueError("No CV data provided")

    personal = data.get('personal')
    if not personal or not isinstance(personal, dict):
        raise ValueError("Missing personal data")
    if not personal.get('name'):
        raise ValueError("Missing candidate name")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    lang = data.get('cv_language', 'en')

    _setup_styles(doc)
    _add_name_header(doc, data['personal'])

    section_order = data.get('section_order',
                             ['summary', 'experience', 'skills', 'projects', 'education', 'courses', 'certifications', 'languages'])

    section_map = {
        'summary': lambda: _add_summary(doc, data.get('summary', ''), lang),
        'experience': lambda: _add_work_experience(doc, data, lang),
        'skills': lambda: _add_skills(doc, data.get('skills', []), lang),
        'projects': lambda: _add_projects(doc, data.get('projects', []), lang),
        'courses': lambda: _add_courses(doc, data.get('courses', []), lang),
        'education': lambda: _add_education(doc, data.get('education', []), lang),
        'languages': lambda: _add_languages(doc, data.get('languages', []), lang),
        'certifications': lambda: _add_certifications(doc, data.get('certifications', []), lang),
    }

    disabled = data.get('disabled_sections', [])
    real_sections = [s for s in section_order
                     if isinstance(s, str) and s in section_map and s not in disabled]

    for idx, section_name in enumerate(real_sections):
        section_map[section_name]()
        if idx < len(real_sections) - 1:
            _add_horizontal_line(doc)

    if data.get('clause_enabled') and data.get('clause_text'):
        _add_horizontal_line(doc)
        _add_clause(doc, data['clause_text'])

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    return tmp.name

def generate_pdf(data):
    """Generate PDF from DOCX using LibreOffice."""
    docx_path = generate_docx(data)
    out_dir = os.path.dirname(docx_path)

    soffice = shutil.which('soffice')
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) not found")

    try:
        lo_profile = tempfile.mkdtemp(prefix='lo_profile_')
        subprocess.run([
            soffice, '--headless',
            f'-env:UserInstallation=file://{lo_profile}',
            '--convert-to', 'pdf',
            '--outdir', out_dir,
            docx_path
        ], check=True, capture_output=True, timeout=60)
        shutil.rmtree(lo_profile, ignore_errors=True)
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
        raise RuntimeError("PDF generation failed")
    finally:
        try:
            os.unlink(docx_path)
        except OSError:
            pass

    pdf_path = docx_path.rsplit('.', 1)[0] + '.pdf'
    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF generation failed - output file not found")
    return pdf_path
